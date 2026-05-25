from __future__ import annotations

import argparse
import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DEFAULT_WORKDIR = Path(__file__).resolve().parent
DEFAULT_ANALYSIS = DEFAULT_WORKDIR / "analysis"
DEFAULT_TEMPLATE = DEFAULT_WORKDIR / "templates" / "我校 2024年度 ESI 高水平论文的产出情况及影响力分析.docx"
DEFAULT_OUTPUT = DEFAULT_WORKDIR / "outputs" / "我校2025年度ESI高水平论文的产出情况及影响力分析.docx"

ANALYSIS = DEFAULT_ANALYSIS
TEMPLATE = DEFAULT_TEMPLATE
OUTPUT = DEFAULT_OUTPUT
OUTPUT_DIR = OUTPUT.parent
SUMMARY: dict[str, object] = {}

SUBJECT_BREAKS = {
    "农业科学": "农业\n科学",
    "生物与生物化学": "生物\n与生\n物化学",
    "化学": "化学",
    "临床医学": "临床\n医学",
    "计算机科学": "计算\n机\n科学",
    "经济学与商学": "经济\n学与\n商学",
    "工程": "工程",
    "环境/生态学": "环境\n/生\n态学",
    "材料科学": "材料\n科学",
    "数学": "数学",
    "精神病学/心理学": "精神\n病学\n/心\n理学",
    "药理学和毒理学": "药理\n学和\n毒理学",
    "植物与动物科学": "植物\n与动\n物\n科学",
    "一般社会科学": "一般社会\n科学",
}

COLLEGE_BREAKS = {
    "财政与税务学院": "财政与\n税务学院",
    "国际经贸学院": "国际经贸\n学院",
    "公共管理学院": "公共管理\n学院",
    "食品科学与工程学院": "食品科学与\n工程学院",
    "计算机与人工智能学院": "计算机与人\n工智能学院",
    "粮食和物资学院": "粮食和物资\n学院",
}

HEATMAP_ANCHORS = {
    1: "FADADE",
    2: "FAD7DB",
    3: "F9D3D7",
    4: "F8D0D4",
    5: "F7CCD0",
    6: "F6C9CD",
    7: "F5C5C9",
    8: "F4C2C5",
    9: "F3BEC2",
    11: "F1B7BB",
    16: "ECA6A9",
    19: "EA9B9E",
    29: "E0787A",
    48: "CF3536",
    63: "C00000",
}

TABLE12_GRID = [960] * 8
TABLE13_GRID = [1416, 458, 458] + [457] * 12 + [470]
TABLE14_GRID = [1656, 852, 1896, 828, 2376, 780]


def pct(count: int, total: int) -> str:
    return f"{round(count / total * 100)}%"


def top_names(counter_map: dict[str, int], limit: int) -> str:
    items = sorted(counter_map.items(), key=lambda kv: (-kv[1], kv[0]))[:limit]
    return "、".join(name for name, _ in items)


def build_intro_paragraph() -> str:
    totals = SUMMARY["totals"]
    dist = SUMMARY["distribution"]
    hot_counts = dist["hot"]
    hot_units = "、".join(name for name, _ in sorted(SUMMARY["hot_college_counts"].items(), key=lambda kv: (-kv[1], kv[0]))[:6])
    hot_extra = (
        f"其中入围2期的热点论文{hot_counts['2']}篇，入围3期的{hot_counts['3']}篇，"
        f"入围5期的{hot_counts['5']}篇，其余{hot_counts['1']}篇入围1期。"
    )
    return (
        f"2025年度，我校共入围{totals['high_level']}篇 ESI 高水平论文，其中高被引论文 {totals['high_cited']}篇，"
        f"热点论文 {totals['hot']}篇。{totals['high_cited']}篇高被引论文中，6 期全部入围的 "
        f"{dist['high_cited']['6']}篇，在本校高水平论文中占比达到 {pct(int(dist['high_cited']['6']), totals['high_level'])}。"
        f"{totals['hot']}篇热点论文主要涉及{hot_units}等院系单位。{hot_extra}"
        f"这些热点论文中，{SUMMARY['hot_also_high_cited']}篇同时为高被引论文或在后续期次进入高被引论文名单。"
    )


def build_subject_paragraph() -> str:
    colleges = SUMMARY["unique_colleges"]
    subjects = SUMMARY["unique_subjects"]
    college_unique = SUMMARY["college_unique_counts"]
    college_subject = SUMMARY["college_subject_counts"]
    subject_totals = defaultdict(int)
    for values in college_subject.values():
        for subject, count in values.items():
            subject_totals[subject] += count

    top_colleges = top_names(college_unique, 4)
    top_subjects = top_names(subject_totals, 4)
    multi_subject_colleges = [
        name
        for name, values in sorted(college_subject.items(), key=lambda kv: (-len(kv[1]), kv[0]))
        if len(values) >= 2
    ]
    sample_college = max(college_subject.items(), key=lambda kv: (len(kv[1]), sum(kv[1].values())))[0]
    sample_subjects = "、".join(
        f"{name}（{count}篇）"
        for name, count in sorted(
            college_subject[sample_college].items(),
            key=lambda kv: (-kv[1], kv[0]),
        )
    )
    source_subjects = sorted(
        SUMMARY["subject_source_colleges"].items(),
        key=lambda kv: (-len(kv[1]), kv[0]),
    )
    source_sentence = (
        f"{source_subjects[0][0]}和{source_subjects[1][0]}的高水平论文均来源于{len(source_subjects[0][1])}个院系单位。"
    )
    return (
        f"我校2025年度入围的ESI高水平论文，来源于全校{len(colleges)}个院系单位，涉及了22个ESI学科中的{len(subjects)}个。"
        f"高水平论文贡献较多的院系为{top_colleges}等。高水平论文主要分布在{top_subjects}等学科。"
        f"{'、'.join(multi_subject_colleges[:7])}等院系贡献了多个 ESI 学科的高水平论文。"
        f"如{sample_college}高水平论文涉及{len(college_subject[sample_college])}个ESI学科，分别为{sample_subjects}。"
        f"{source_sentence}"
    )


def build_author_paragraph() -> str:
    author_labels = {name for authors in SUMMARY["author_counts"].values() for name in authors}
    author_total = len(author_labels)
    pending_total = len(SUMMARY.get("pinyin_pending_papers", []))
    unresolved_total = len(SUMMARY.get("unresolved_author_papers", []))
    sentence = (
        f"我校 2025年度入围的 {SUMMARY['totals']['high_level']}篇高水平论文，"
        f"共涉及{len(SUMMARY['unique_colleges'])}个院系单位的{author_total}位作者。"
    )
    if pending_total:
        sentence += f"其中{pending_total}篇论文在作者清洗后仍保留了部分或全部拼音作者；在表14代表作者统计层面，仍有{unresolved_total}篇论文暂按附表拼音列示。"
    return sentence


def build_author_note() -> str:
    note = (
        "（注：所属院系以发文地址为准；部分论文为两个及以上院系共同发文。"
        "报告中所涉及的高水平论文作者统计，依据附表中可识别的本校作者信息进行归并，"
        "同一篇论文在同一院系内记1位代表作者。"
    )
    pending_total = len(SUMMARY.get("pinyin_pending_papers", []))
    unresolved_total = len(SUMMARY.get("unresolved_author_papers", []))
    if pending_total:
        note += f"另有{pending_total}篇论文仍保留部分或全部拼音作者，其中{unresolved_total}篇论文在表14中暂按附表拼音列示。"
    note += "）"
    return note


def build_author_correction_note() -> str:
    pending = SUMMARY.get("pinyin_pending_papers", [])
    if not pending:
        return ""
    parts = []
    for item in pending:
        short_title = item["title"]
        if len(short_title) > 42:
            short_title = short_title[:42] + "..."
        resolved = "；".join(item.get("resolved_authors", [])) or "无"
        pending_names = "；".join(item.get("display_authors", [])) or "无"
        parts.append(f"《{short_title}》已识别作者为{resolved}，待核作者为{pending_names}")
    return "作者修正说明：本次复核后，年度总表中仍有3篇论文保留部分或全部拼音作者，具体为：" + "；".join(parts) + "。其中表14代表作者统计层面，暂按拼音列示的为 Chen,Tianyu。"


def remove_run(run) -> None:
    run._element.getparent().remove(run._element)


def clear_paragraph_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        remove_run(run)


def set_run_fonts(
    run,
    *,
    ascii_font: str | None = None,
    east_asia_font: str | None = None,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
) -> None:
    if ascii_font is not None:
        run.font.name = ascii_font
    if east_asia_font is not None or ascii_font is not None:
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        if ascii_font is not None:
            r_fonts.set(qn("w:ascii"), ascii_font)
            r_fonts.set(qn("w:hAnsi"), ascii_font)
            r_fonts.set(qn("w:cs"), ascii_font)
        if east_asia_font is not None:
            r_fonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def replace_paragraph_keep_style(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in list(paragraph.runs[1:]):
            remove_run(run)
        return
    paragraph.add_run(text)


def replace_paragraph_with_font(
    paragraph,
    text: str,
    *,
    ascii_font: str,
    east_asia_font: str,
    size: float,
) -> None:
    clear_paragraph_runs(paragraph)
    run = paragraph.add_run(text)
    set_run_fonts(run, ascii_font=ascii_font, east_asia_font=east_asia_font, size=size)


def set_table_width(table, width: int, width_type: str = "dxa") -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), width_type)
    tbl_w.set(qn("w:w"), str(width))


def set_table_alignment(table, value: str = "center") -> None:
    tbl_pr = table._tbl.tblPr
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), value)


def set_table_indent(table, value: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(value))


def set_table_grid(table, widths: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_table_borders(table, size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        child = borders.find(qn(f"w:{edge}"))
        if child is None:
            child = OxmlElement(f"w:{edge}")
            borders.append(child)
        child.set(qn("w:val"), "single")
        child.set(qn("w:sz"), size)
        child.set(qn("w:space"), "0")
        child.set(qn("w:color"), "000000")


def set_row_height(row, value: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_h = tr_pr.find(qn("w:trHeight"))
    if tr_h is None:
        tr_h = OxmlElement("w:trHeight")
        tr_pr.append(tr_h)
    tr_h.set(qn("w:val"), str(value))


def set_cell_fill(cell, fill: str | None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if fill is None:
        if shd is not None:
            tc_pr.remove(shd)
        return
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_text(
    cell,
    text: str,
    *,
    ascii_font: str,
    east_asia_font: str,
    size: float = 12.0,
    color: str = "000000",
    bold: bool = False,
    fill: str | None = None,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    clear_paragraph_runs(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_fonts(
        run,
        ascii_font=ascii_font,
        east_asia_font=east_asia_font,
        size=size,
        color=color,
        bold=bold,
    )
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_fill(cell, fill)


def rgb_triplet(hex_color: str) -> tuple[int, int, int]:
    return (
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )


def interpolated_fill(count: int) -> str | None:
    if count <= 0:
        return None
    if count in HEATMAP_ANCHORS:
        return HEATMAP_ANCHORS[count]
    anchor_counts = sorted(HEATMAP_ANCHORS)
    if count <= anchor_counts[0]:
        return HEATMAP_ANCHORS[anchor_counts[0]]
    if count >= anchor_counts[-1]:
        return HEATMAP_ANCHORS[anchor_counts[-1]]
    lower = max(anchor for anchor in anchor_counts if anchor < count)
    upper = min(anchor for anchor in anchor_counts if anchor > count)
    low_rgb = rgb_triplet(HEATMAP_ANCHORS[lower])
    high_rgb = rgb_triplet(HEATMAP_ANCHORS[upper])
    ratio = (count - lower) / (upper - lower)
    blended = tuple(round(low + (high - low) * ratio) for low, high in zip(low_rgb, high_rgb))
    return "".join(f"{value:02X}" for value in blended)


def display_college(name: str) -> str:
    return COLLEGE_BREAKS.get(name, name)


def build_distribution_table(doc: Document):
    total = SUMMARY["totals"]["high_level"]
    dist = SUMMARY["distribution"]
    table = doc.add_table(rows=5, cols=8)
    table.style = "Normal Table"
    set_table_width(table, 7680)
    set_table_alignment(table)
    set_table_grid(table, TABLE12_GRID)
    set_table_borders(table)
    for row in table.rows:
        set_row_height(row, 288)

    headers = ["入围期数", "入围1期", "入围2期", "入围3期", "入围4期", "入围5期", "入围6期", "合计"]
    for idx, text in enumerate(headers):
        set_cell_text(table.cell(0, idx), text, ascii_font="宋体", east_asia_font="宋体", bold=True)

    high_cited_label = table.cell(1, 0).merge(table.cell(2, 0))
    hot_label = table.cell(3, 0).merge(table.cell(4, 0))
    set_cell_text(high_cited_label, "高被引论文数量及占比", ascii_font="宋体", east_asia_font="宋体")
    set_cell_text(hot_label, "热点论文数量及占比", ascii_font="宋体", east_asia_font="宋体")

    hc = dist["high_cited"]
    hot = dist["hot"]
    rows = [
        [hc["1"], hc["2"], hc["3"], hc["4"], hc["5"], hc["6"], SUMMARY["totals"]["high_cited"]],
        [
            pct(int(hc["1"]), total),
            pct(int(hc["2"]), total),
            pct(int(hc["3"]), total),
            pct(int(hc["4"]), total),
            pct(int(hc["5"]), total),
            pct(int(hc["6"]), total),
            pct(SUMMARY["totals"]["high_cited"], total),
        ],
        [hot["1"], hot["2"], hot["3"], hot["4"], hot["5"], hot["6"], SUMMARY["totals"]["hot"]],
        [
            pct(int(hot["1"]), total),
            pct(int(hot["2"]), total),
            pct(int(hot["3"]), total),
            pct(int(hot["4"]), total),
            pct(int(hot["5"]), total),
            pct(int(hot["6"]), total),
            pct(SUMMARY["totals"]["hot"], total),
        ],
    ]
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row, start=1):
            set_cell_text(
                table.cell(row_idx, col_idx),
                str(value),
                ascii_font="Times New Roman",
                east_asia_font="宋体",
            )
    return table


def build_subject_table(doc: Document):
    subjects = SUMMARY["unique_subjects"]
    colleges = SUMMARY["unique_colleges"]
    table = doc.add_table(rows=len(colleges) + 1, cols=len(subjects) + 2)
    table.style = "Normal Table"
    set_table_width(table, 5000, width_type="pct")
    set_table_indent(table, -294)
    set_table_grid(table, TABLE13_GRID)
    set_table_borders(table)
    set_row_height(table.rows[0], 1618)

    headers = ["学院/\n学科"] + [SUBJECT_BREAKS.get(subject, subject) for subject in subjects] + ["总计"]
    for idx, text in enumerate(headers):
        set_cell_text(table.cell(0, idx), text, ascii_font="宋体", east_asia_font="宋体", bold=True)

    college_subject = SUMMARY["college_subject_counts"]
    for row_idx, college in enumerate(colleges, start=1):
        set_cell_text(table.cell(row_idx, 0), display_college(college), ascii_font="宋体", east_asia_font="宋体")
        row_total = 0
        for col_idx, subject in enumerate(subjects, start=1):
            count = int(college_subject.get(college, {}).get(subject, 0))
            row_total += count
            set_cell_text(
                table.cell(row_idx, col_idx),
                "" if count == 0 else str(count),
                ascii_font="Times New Roman",
                east_asia_font="宋体",
                fill=interpolated_fill(count),
            )
        set_cell_text(
            table.cell(row_idx, len(subjects) + 1),
            str(row_total),
            ascii_font="Times New Roman",
            east_asia_font="宋体",
            fill=interpolated_fill(row_total),
        )
    return table


def build_author_columns():
    college_unique = SUMMARY["college_unique_counts"]
    author_counts = SUMMARY["author_counts"]
    blocks = []
    for college in SUMMARY["unique_colleges"]:
        authors = sorted(author_counts.get(college, {}).items(), key=lambda kv: (-kv[1], kv[0]))
        block = [(college, str(college_unique[college]), True)]
        for author, count in authors:
            block.append((author, str(count), False))
        blocks.append(block)

    sizes = [len(block) for block in blocks]
    prefix = [0]
    for size in sizes:
        prefix.append(prefix[-1] + size)

    best = None
    best_splits = (0, 0)
    n = len(blocks)
    for i in range(1, n - 1):
        for j in range(i + 1, n):
            parts = [
                prefix[i] - prefix[0],
                prefix[j] - prefix[i],
                prefix[n] - prefix[j],
            ]
            score = max(parts)
            if best is None or score < best:
                best = score
                best_splits = (i, j)

    i, j = best_splits
    grouped = [blocks[:i], blocks[i:j], blocks[j:]]
    columns = []
    for group in grouped:
        column = []
        for block in group:
            column.extend(block)
        columns.append(column)

    max_height = max(len(column) for column in columns)
    for column in columns:
        while len(column) < max_height:
            column.append(("", "", False))
    return columns, max_height


def build_author_table(doc: Document):
    columns, max_height = build_author_columns()
    table = doc.add_table(rows=max_height + 1, cols=6)
    table.style = "Normal Table"
    set_table_width(table, 8388)
    set_table_alignment(table)
    set_table_grid(table, TABLE14_GRID)
    set_table_borders(table)
    set_row_height(table.rows[0], 288)

    headers = ["学院/作者", "数量", "学院/作者", "数量", "学院/作者", "数量"]
    for idx, text in enumerate(headers):
        set_cell_text(table.cell(0, idx), text, ascii_font="宋体", east_asia_font="宋体", bold=True)

    for row_idx in range(max_height):
        for group_idx in range(3):
            label, count, is_college = columns[group_idx][row_idx]
            color = "FF0000" if is_college else "000000"
            text = display_college(label) if is_college else label
            set_cell_text(
                table.cell(row_idx + 1, group_idx * 2),
                text,
                ascii_font="宋体",
                east_asia_font="宋体",
                color=color,
            )
            set_cell_text(
                table.cell(row_idx + 1, group_idx * 2 + 1),
                count,
                ascii_font="Times New Roman",
                east_asia_font="宋体",
                color=color,
            )
    return table


def replace_table(old_table, new_table) -> None:
    old_tbl = old_table._tbl
    new_tbl = new_table._tbl
    old_tbl.addprevious(new_tbl)
    old_tbl.getparent().remove(old_tbl)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="根据 analysis 结果生成年度正式报告 DOCX")
    parser.add_argument("--analysis-dir", type=Path, default=DEFAULT_ANALYSIS, help="analysis 目录")
    parser.add_argument("--template-docx", type=Path, default=DEFAULT_TEMPLATE, help="2024 年报告模板 DOCX")
    parser.add_argument("--output-docx", type=Path, default=DEFAULT_OUTPUT, help="输出 DOCX 路径")
    return parser.parse_args()


def generate_report_doc(analysis_dir: Path, template_docx: Path, output_docx: Path) -> Path:
    global ANALYSIS, TEMPLATE, OUTPUT, OUTPUT_DIR, SUMMARY

    ANALYSIS = analysis_dir
    TEMPLATE = template_docx
    OUTPUT = output_docx
    OUTPUT_DIR = OUTPUT.parent
    SUMMARY = json.loads((ANALYSIS / "summary.json").read_text(encoding="utf-8"))

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(TEMPLATE)

    paragraphs = doc.paragraphs
    replace_paragraph_keep_style(paragraphs[1], "6 我校 2025年度 ESI 高水平论文的产出情况及影响力分析")
    replace_paragraph_keep_style(paragraphs[2], "6.1 我校 2025年度 ESI 高水平论文入围情况 ")
    replace_paragraph_with_font(paragraphs[3], build_intro_paragraph(), ascii_font="Times New Roman", east_asia_font="仿宋", size=15)
    replace_paragraph_keep_style(paragraphs[4], "表12我校2025年度ESI高水平论文入围期数分布及占比")

    replace_paragraph_keep_style(paragraphs[5], "6.2 我校 2025年度 ESI 高水平论文的院系及学科分布")
    replace_paragraph_keep_style(paragraphs[6], "表13 2025年入围的ESI高水平论文所属单位及学科分布")
    replace_paragraph_keep_style(paragraphs[7], "（注：所属院系以发文地址为准；部分论文为两个及以上院系共同发文。）")
    replace_paragraph_with_font(paragraphs[8], build_subject_paragraph(), ascii_font="Times New Roman", east_asia_font="仿宋", size=15)

    replace_paragraph_keep_style(paragraphs[9], "6.3 我校2025年度ESI高水平论文的院系及作者分布 ")
    replace_paragraph_with_font(paragraphs[10], build_author_paragraph(), ascii_font="Times New Roman", east_asia_font="仿宋", size=15)
    replace_paragraph_keep_style(paragraphs[11], "表14我校 ESI 高水平论文的院系及作者分布")
    replace_paragraph_keep_style(paragraphs[12], build_author_note())
    old_table12, old_table13, old_table14 = list(doc.tables)
    replace_table(old_table12, build_distribution_table(doc))
    replace_table(old_table13, build_subject_table(doc))
    replace_table(old_table14, build_author_table(doc))

    if doc.paragraphs and not doc.paragraphs[-1].text:
        last_paragraph = doc.paragraphs[-1]._element
        last_paragraph.getparent().remove(last_paragraph)

    correction_note = build_author_correction_note()
    if correction_note:
        paragraph = doc.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.line_spacing = 1.5
        run = paragraph.add_run(correction_note)
        set_run_fonts(run, ascii_font="Times New Roman", east_asia_font="仿宋", size=12)

    doc.save(OUTPUT)
    return OUTPUT


def main() -> None:
    args = parse_args()
    output = generate_report_doc(args.analysis_dir, args.template_docx, args.output_docx)
    print(output)


if __name__ == "__main__":
    main()
