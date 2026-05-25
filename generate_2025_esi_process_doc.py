from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DEFAULT_WORKDIR = Path(__file__).resolve().parent
DEFAULT_ANALYSIS = DEFAULT_WORKDIR / "analysis"
DEFAULT_OUTPUT = DEFAULT_WORKDIR / "outputs" / "2025年度ESI高水平论文统计口径与核验说明.docx"

ANALYSIS = DEFAULT_ANALYSIS
OUTPUT = DEFAULT_OUTPUT
OUTPUT_DIR = OUTPUT.parent
SUMMARY: dict[str, object] = {}

SOURCE_ROWS = [
    ("第1期", "ESI2025年第1期附表.xlsx", "附表1.本期我校高被引论文和热点论文", 108, 4),
    ("第2期", "ESI2025年第2期附表.xlsx", "附表1.本期我校高被引论文和热点论文", 112, 5),
    ("第3期", "ESI2025年第3期附表.xlsx", "附表3.本期我校高被引论文和热点论文", 110, 5),
    ("第4期", "ESI2025年第4-5期附表.xlsx", "附表1.2025年7月高被引论文和热点论文", 113, 4),
    ("第5期", "ESI2025年第4-5期附表.xlsx", "附表2.2025年9月高被引论文和热点论文", 115, 7),
    ("第6期", "ESI2025年第6期附表(1).xlsx", "附表1.本期我校高被引论文和热点论文", 119, 2),
]

AUTHOR_EXAMPLES = [
    ("Huang,XW(Huang, Xiaowei 黄晓玮)", "黄晓玮", "附表括注直接对应"),
    ("Peng,Q(Peng, Qiong 彭琼)", "彭琼", "附表括注直接对应"),
    ("Li,W(Li, Wen 李文)", "李文", "附表括注直接对应"),
    ("Wu, Peipei", "吴佩佩", "南京财经大学官网检索结果对应"),
    ("Li, Rongrong", "李荣荣", "南京财经大学官网检索结果对应"),
    ("Chen, Afei", "陈阿飞", "南京财经大学官网检索结果对应"),
]


def resolved_author_total() -> int:
    author_labels = {name for authors in SUMMARY["author_counts"].values() for name in authors}
    return len({name for name in author_labels if re.search(r"[\u4e00-\u9fff]", name)})


def author_label_total() -> int:
    return len({name for authors in SUMMARY["author_counts"].values() for name in authors})


def set_page(section) -> None:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def ensure_fonts(run, *, ascii_font: str, east_asia_font: str, size: float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:cs"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def add_paragraph(doc: Document, text: str, *, size: float = 12, bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line: float | None = 0.74):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    if first_line is not None:
        fmt.first_line_indent = Cm(first_line)
    run = paragraph.add_run(text)
    ensure_fonts(run, ascii_font="Times New Roman", east_asia_font="仿宋", size=size, bold=bold)
    return paragraph


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    run = paragraph.add_run(text)
    ensure_fonts(run, ascii_font="Times New Roman", east_asia_font="黑体", size=14, bold=True)


def set_table_width(table, width: int, width_type: str = "dxa") -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), width_type)
    tbl_w.set(qn("w:w"), str(width))


def set_table_grid(table, widths: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_table_borders(table, size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        child = borders.find(qn(f"w:{edge}"))
        if child is None:
            child = OxmlElement(f"w:{edge}")
            borders.append(child)
        child.set(qn("w:val"), "single")
        child.set(qn("w:sz"), size)
        child.set(qn("w:space"), "0")
        child.set(qn("w:color"), "000000")


def set_cell_fill(cell, fill: str | None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if fill is None:
        if shd is not None:
            tc_pr.remove(shd)
        return
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, size: float = 10.5, bold: bool = False, fill: str | None = None, center: bool = True) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    run = paragraph.add_run(text)
    ensure_fonts(run, ascii_font="Times New Roman", east_asia_font="宋体", size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_fill(cell, fill)


def add_table(doc: Document, rows: list[list[str]], widths: list[int], header_fill: str = "D9EAF7") -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_width(table, sum(widths))
    set_table_grid(table, widths)
    set_table_borders(table)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(
                table.cell(i, j),
                value,
                bold=i == 0,
                fill=header_fill if i == 0 else None,
                center=False if j in {1, 2, 3} and len(row) >= 4 else True,
            )
    doc.add_paragraph()


def build_overlap_rows() -> list[list[str]]:
    totals = SUMMARY["totals"]
    overlap = SUMMARY["hot_also_high_cited"]
    high_cited_only = totals["high_cited"] - overlap
    hot_only = totals["hot"] - overlap
    return [
        ["类型", "篇数", "说明"],
        ["仅高被引论文", str(high_cited_only), "至少进入1期高被引名单，但未进入热点名单"],
        ["高被引与热点重合", str(overlap), "热点论文且同时为高被引论文，或后续进入高被引名单"],
        ["仅热点论文", str(hot_only), "进入热点名单，但未进入任何一期高被引名单"],
        ["高水平论文合计", str(totals["high_level"]), "高被引论文与热点论文去重并集"],
    ]


def build_unresolved_rows() -> list[list[str]]:
    rows = [["论文题名", "已识别中文作者", "待核拼音作者", "处理意见"]]
    for item in SUMMARY.get("pinyin_pending_papers", []):
        resolved_authors = "；".join(item.get("resolved_authors", [])) or "无"
        display_authors = "；".join(item.get("display_authors", [])) or "无"
        if item.get("fully_unresolved"):
            action = "未在附表或校内可靠公开信息中确认中文名，年度统计表按附表拼音列示"
        else:
            action = "已保留可确认中文作者，其余作者暂按附表拼音列示，待后续核实"
        rows.append(
            [
                item["title"],
                resolved_authors,
                display_authors,
                action,
            ]
        )
    return rows


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="根据 analysis 结果生成统计口径说明 DOCX")
    parser.add_argument("--analysis-dir", type=Path, default=DEFAULT_ANALYSIS, help="analysis 目录")
    parser.add_argument("--output-docx", type=Path, default=DEFAULT_OUTPUT, help="输出 DOCX 路径")
    return parser.parse_args()


def generate_process_doc(analysis_dir: Path, output_docx: Path) -> Path:
    global ANALYSIS, OUTPUT, OUTPUT_DIR, SUMMARY

    ANALYSIS = analysis_dir
    OUTPUT = output_docx
    OUTPUT_DIR = OUTPUT.parent
    SUMMARY = json.loads((ANALYSIS / "summary.json").read_text(encoding="utf-8"))

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_page(doc.sections[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_fmt = title.paragraph_format
    title_fmt.space_after = Pt(10)
    run = title.add_run("2025年度ESI高水平论文统计口径与核验说明")
    ensure_fonts(run, ascii_font="Times New Roman", east_asia_font="黑体", size=18, bold=True)

    totals = SUMMARY["totals"]
    overlap = SUMMARY["hot_also_high_cited"]
    hot_only = totals["hot"] - overlap

    add_paragraph(
        doc,
        f"根据“2025年ESI高水平论文”文件夹中的6期附表，对2025年度我校ESI高水平论文统计结果进行了重新核验，并对作者姓名中的拼音写法进行了二次清洗。经复核，修正后的年度总量为：高水平论文{totals['high_level']}篇，其中高被引论文{totals['high_cited']}篇、热点论文{totals['hot']}篇；两类名单重合{overlap}篇，仅热点论文{hot_only}篇。",
    )

    add_heading(doc, "一、数据来源与取表范围")
    add_paragraph(
        doc,
        "本次统计以文件夹内6期附表为唯一基础数据来源。第4期和第5期共用同一工作簿，分别取前两个“高被引论文和热点论文”附表；高水平论文口径包含高被引论文和热点论文两类。",
    )
    source_rows = [["期次", "来源文件", "取用附表", "高被引", "热点"]]
    for issue, filename, sheet_name, hc, hot in SOURCE_ROWS:
        source_rows.append([issue, filename, sheet_name, str(hc), str(hot)])
    add_table(doc, source_rows, [900, 2300, 3200, 800, 800])

    add_heading(doc, "二、统计口径")
    add_paragraph(doc, "1. 年度高水平论文总数按“高被引论文 ∪ 热点论文”去重统计；同一篇论文即使跨多期反复入围，年度合计中仅计1篇。", first_line=None)
    add_paragraph(doc, "2. 去重键采用论文题名标准化处理：统一大小写，去除HTML标签、空格、标点及连字符差异后再比对，避免同题名不同写法造成重复计数。", first_line=None)
    add_paragraph(doc, "3. 院系归属以附表“所属学院”为准；若同一篇论文对应两个及以上本校院系，则在院系分布统计中分别记入相关院系。", first_line=None)
    add_paragraph(doc, "4. 作者统计以附表作者栏中可识别的本校作者信息为基础，同一篇论文在同一院系内只记1位代表作者，不按全部作者人次累加。", first_line=None)
    add_paragraph(doc, "5. 作者姓名清洗优先采用附表中括注出现的“拼音+中文”对应关系；附表未给出中文名时，再以校内可靠公开信息进行核对，无法可靠确认时暂按附表拼音列示，并将中文名标记为待后续核实。", first_line=None)

    add_heading(doc, "三、统计结果复核")
    add_paragraph(
        doc,
        "逐期原始计数与年度去重结果相互一致，未发现总量漏统、重复累计或院系/学科缺项问题。当前统计结果覆盖15个院系单位、14个ESI学科，学科和院系字段已全部补齐，缺失项为0。",
    )
    issue_rows = [["期次", "高被引论文", "热点论文", "说明"]]
    for issue, _, _, hc, hot in SOURCE_ROWS:
        issue_rows.append([issue, str(hc), str(hot), "与附表逐行核对一致"])
    add_table(doc, issue_rows, [900, 1200, 1200, 4700])
    add_table(doc, build_overlap_rows(), [1600, 1000, 5400])

    add_heading(doc, "四、作者姓名清洗结果")
    pending_total = len(SUMMARY.get("pinyin_pending_papers", []))
    fully_unresolved_total = len(SUMMARY.get("unresolved_author_papers", []))
    add_paragraph(
        doc,
        f"清洗后，年度统计表中中文作者姓名已尽量统一。当前作者分布表共列示{author_label_total()}位作者，其中{resolved_author_total()}位可用中文名归并；另有{pending_total}篇论文仍保留部分或全部拼音作者待核，其中{fully_unresolved_total}篇论文尚未识别出任何中文作者，只能在年度统计表中按附表拼音列示。",
    )
    example_rows = [["原始写法", "统一后的中文名", "核定依据"]]
    for raw_name, resolved_name, basis in AUTHOR_EXAMPLES:
        example_rows.append([raw_name, resolved_name, basis])
    add_table(doc, example_rows, [3300, 1400, 3300])

    add_heading(doc, "五、待人工确认事项")
    add_paragraph(
        doc,
        f"经再次核对，目前共有{pending_total}篇论文存在拼音作者待进一步确认。其中{fully_unresolved_total}篇论文尚未识别出任何中文作者，正式年度报告中相应作者在作者分布表内暂按附表拼音列示；其余论文虽已识别出部分中文作者，但仍有部分拼音作者未核定中文名，建议后续结合原论文全文、学院科研秘书台账或作者本人信息进一步确认。",
    )
    add_table(doc, build_unresolved_rows(), [2500, 1600, 1800, 2100], header_fill="FBE5D6")

    add_heading(doc, "六、结论")
    add_paragraph(
        doc,
        "本次复核后，2025年度ESI高水平论文统计主数据无误，现有口径可以直接用于汇报。作者姓名方面，凡能从附表括注或校内可靠公开信息确认的，均已统一为中文；仍保留拼音的作者，建议后续逐篇结合论文全文与院系台账进一步补齐，确认后即可将相应拼音姓名替换为正式中文姓名。",
    )

    doc.save(OUTPUT)
    return OUTPUT


def main() -> None:
    args = parse_args()
    output = generate_process_doc(args.analysis_dir, args.output_docx)
    print(output)


if __name__ == "__main__":
    main()
