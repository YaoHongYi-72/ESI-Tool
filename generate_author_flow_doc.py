from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DEFAULT_WORKDIR = Path(__file__).resolve().parent
DEFAULT_ANALYSIS = DEFAULT_WORKDIR / "analysis"
DEFAULT_OUTPUT = DEFAULT_WORKDIR / "outputs" / "2025年度ESI高水平论文作者统计流程详解.docx"

ANALYSIS = DEFAULT_ANALYSIS
OUTPUT = DEFAULT_OUTPUT
OUTPUT_DIR = OUTPUT.parent
SUMMARY: dict[str, object] = {}


def author_label_total() -> int:
    return len({name for authors in SUMMARY["author_counts"].values() for name in authors})


def resolved_author_total() -> int:
    labels = {name for authors in SUMMARY["author_counts"].values() for name in authors}
    return len({name for name in labels if re.search(r"[\u4e00-\u9fff]", name)})


def set_page(section) -> None:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def ensure_fonts(run, *, ascii_font: str, east_asia_font: str, size: float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:cs"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.space_after = Pt(10)
    run = p.add_run(text)
    ensure_fonts(run, ascii_font="Times New Roman", east_asia_font="黑体", size=18, bold=True)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    run = p.add_run(text)
    ensure_fonts(run, ascii_font="Times New Roman", east_asia_font="黑体", size=14, bold=True)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: float = 12,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line: float | None = 0.74,
    east_asia_font: str = "仿宋",
    ascii_font: str = "Times New Roman",
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    if first_line is not None:
        fmt.first_line_indent = Cm(first_line)
    run = p.add_run(text)
    ensure_fonts(run, ascii_font=ascii_font, east_asia_font=east_asia_font, size=size, bold=bold)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt = p.paragraph_format
        fmt.left_indent = Cm(0.8)
        fmt.right_indent = Cm(0.4)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        run = p.add_run(line)
        ensure_fonts(run, ascii_font="Courier New", east_asia_font="等线", size=9.5)
    doc.add_paragraph()


def set_table_width(table, width: int, width_type: str = "dxa") -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), width_type)
    tbl_w.set(qn("w:w"), str(width))


def set_table_grid(table, widths: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_table_borders(table, size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        child = borders.find(qn(f"w:{edge}"))
        if child is None:
            child = OxmlElement(f"w:{edge}")
            borders.append(child)
        child.set(qn("w:val"), "single")
        child.set(qn("w:sz"), size)
        child.set(qn("w:space"), "0")
        child.set(qn("w:color"), "000000")


def set_cell_fill(cell, fill: str | None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if fill is None:
        if shd is not None:
            tc_pr.remove(shd)
        return
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, size: float = 10.5, bold: bool = False, fill: str | None = None, center: bool = True) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    run = paragraph.add_run(text)
    ensure_fonts(run, ascii_font="Times New Roman", east_asia_font="宋体", size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_fill(cell, fill)


def add_table(doc: Document, rows: list[list[str]], widths: list[int], header_fill: str = "D9EAF7") -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_width(table, sum(widths))
    set_table_grid(table, widths)
    set_table_borders(table)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(
                table.cell(i, j),
                value,
                bold=i == 0,
                fill=header_fill if i == 0 else None,
                center=False if j > 0 else True,
            )
    doc.add_paragraph()


def build_core_status_rows() -> list[list[str]]:
    overlap = SUMMARY["hot_also_high_cited"]
    return [
        ["指标", "当前结果", "说明"],
        ["年度高水平论文总数", str(SUMMARY["totals"]["high_level"]), "高被引论文与热点论文去重并集"],
        ["高被引论文", str(SUMMARY["totals"]["high_cited"]), "按题名去重后的年度高被引论文数"],
        ["热点论文", str(SUMMARY["totals"]["hot"]), "按题名去重后的年度热点论文数"],
        ["两类重合论文", str(overlap), "同时属于热点论文和高被引论文，或后续进入高被引名单"],
        ["表14代表作者标签数", str(author_label_total()), "作者分布表中实际出现的作者标签总数"],
        ["其中可用中文名归并", str(resolved_author_total()), "在表14中可直接用中文名展示的作者标签数"],
        ["仍含拼音待核的论文", str(len(SUMMARY.get("pinyin_pending_papers", []))), "总表作者列仍保留部分或全部拼音作者的论文篇数"],
        ["完全未识别中文作者", str(len(SUMMARY.get("unresolved_author_papers", []))), "该类论文在表14中只能按拼音列示代表作者"],
    ]


def build_object_rows() -> list[list[str]]:
    return [
        ["对象", "创建时机", "关键字段", "作用"],
        ["PaperRecord", "逐行读取每一期 Excel 时创建", "期次、类别、原始题名、原始作者串、学院、学科、作者拆分结果", "表示附表中的一条原始论文记录"],
        ["PaperAggregate", "按标准化题名聚合时创建", "高被引期次集合、热点期次集合、学院集合、学科集合、中文作者、待核拼音作者", "表示去重后的一篇年度论文"],
    ]


def build_pending_rows() -> list[list[str]]:
    rows = [["论文题名", "已识别中文作者", "待核拼音作者", "当前说明"]]
    for item in SUMMARY.get("pinyin_pending_papers", []):
        rows.append(
            [
                item["title"],
                "；".join(item.get("resolved_authors", [])) or "无",
                "；".join(item.get("display_authors", [])) or "无",
                "整篇仍待核" if item.get("fully_unresolved") else "已保留中文作者，其余作者暂以拼音保留",
            ]
        )
    return rows


def build_examples_rows() -> list[list[str]]:
    return [
        ["论文", "总表作者列写法", "表14如何计入", "说明"],
        [
            "GOVERNMENT INVESTMENT, HUMAN CAPITAL FLOW, AND URBAN INNOVATION...",
            "贺伊琦；Li,Zhihui；Wang,Xinyue",
            "会计学院记给贺伊琦",
            "总表展示全部已识别作者；表14只在该学院下选1位代表作者",
        ],
        [
            "AN INTELLIGENT THYMOL/ALIZARIN-LOADED...",
            "彭琼；李文；林亲录；Deng,Jing",
            "食品科学与工程学院记给彭琼",
            "中文作者已统一，剩余拼音作者继续保留待核",
        ],
        [
            "MEDIA ATTENTION, INFORMATION ASYMMETRY AND AGRIBUSINESS ESG RATING DIVERGENCE",
            "Chen,Tianyu；Yin,Ding；Li,Xinghua；Xu,Xuegao",
            "金融学院表14暂以 Chen,Tianyu 记1次",
            "整篇未能可靠核出中文名，因此代表作者也按拼音列示",
        ],
    ]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="生成作者统计流程详解 DOCX")
    parser.add_argument("--analysis-dir", type=Path, default=DEFAULT_ANALYSIS, help="analysis 目录")
    parser.add_argument("--output-docx", type=Path, default=DEFAULT_OUTPUT, help="输出 DOCX 路径")
    return parser.parse_args()


def generate_author_flow_doc(analysis_dir: Path, output_docx: Path) -> Path:
    global ANALYSIS, OUTPUT, OUTPUT_DIR, SUMMARY

    ANALYSIS = analysis_dir
    OUTPUT = output_docx
    OUTPUT_DIR = OUTPUT.parent
    SUMMARY = json.loads((ANALYSIS / "summary.json").read_text(encoding="utf-8"))

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_page(doc.sections[0])

    add_title(doc, "2025年度ESI高水平论文作者统计流程详解")
    add_paragraph(
        doc,
        "本文专门说明“作者”这一部分是怎么统计出来的，重点解释：原始作者串如何拆分、拼音作者如何尽量回填为中文、去重后的一篇论文在循环里发生了什么、表14为什么每个学院只记1位代表作者，以及当前仍有哪些作者信息需要人工继续核实。",
    )

    add_heading(doc, "一、当前作者统计口径与结果")
    add_paragraph(
        doc,
        "目前作者统计同时服务两个输出口径：一是“去重总表”的本校作者列，用于尽量完整展示每篇论文已识别的本校作者；二是正式年度报告中的表14，用于按“学院 × 代表作者”做分布统计。前者可以在一篇论文里展示多位作者，后者则坚持“同一篇论文在同一学院内只记1位代表作者”。",
        first_line=None,
    )
    add_table(doc, build_core_status_rows(), [1700, 1200, 4900])

    add_heading(doc, "二、输入数据与两个核心对象")
    add_paragraph(
        doc,
        "作者统计并不是直接从最终 Excel 一步得出的，而是先把6期附表逐行读成原始记录，再按题名聚合成年度论文对象。作者信息也随着这个过程从“单期、单行”的原始状态，逐步变成“年度去重后”的汇总状态。",
    )
    add_table(doc, build_object_rows(), [1200, 1800, 2200, 2760])

    add_heading(doc, "三、作者统计的详细流程")
    add_paragraph(doc, "1. 逐期读取附表时，程序会先创建 PaperRecord。每条记录至少保留：第几期、是高被引还是热点、论文题名、作者原始字符串、所属学院、所属ESI学科。此时作者还没有被统一成中文名。", first_line=None)
    add_paragraph(doc, "2. 对每条 PaperRecord，先执行作者字符串拆分。程序以分号为界把作者栏拆成多个 token，再分别提取两类信息：一类是 token 里直接出现的中文姓名，另一类是可用于回填的拼音别名。例如 “Xu,X(Xu, Xin 徐鑫)” 会同时得到中文名“徐鑫”和拼音别名 “xuxin”。", first_line=None)
    add_paragraph(doc, "3. 程序随后构建全年的作者别名映射表。做法是：凡是一条记录里同时出现了拼音别名和中文名，就把两者建立对应关系；之后再叠加少量人工确认过的别名，例如吴佩佩、李荣荣、陈阿飞。这样，后面只写拼音的作者就有机会被同年其他记录“带出来”。", first_line=None)
    add_paragraph(doc, "4. 在 resolve_record_authors 这一步，每条 PaperRecord 会生成两份结果：一份是 chinese_names，表示已经成功落成中文的本校作者；另一份是 unresolved_author_tokens，表示当前仍只剩拼音、暂时不能安全翻译的人。处理顺序是：先收直接出现的中文名，再尝试用别名映射回填，最后才把仍无法确认的 token 放进待核列表。", first_line=None)
    add_paragraph(doc, "5. 当所有 PaperRecord 都完成上述处理后，程序开始按标准化题名聚合，创建 PaperAggregate。这里才是真正的“去重”阶段：同一篇论文不论在哪几期重复出现，最终都只对应1个 PaperAggregate。作者信息在这里会被合并，具体表现为：中文作者列表去重保留，仍待核的拼音 token 也去重保留。", first_line=None)
    add_paragraph(doc, "6. 对聚合后的论文，程序会进一步反推出“作者 -> 学院”映射。反推依据主要来自两块：一是2024年文档里已有的作者-学院种子表；二是2025年那些只对应单一学院的论文。因为单学院论文更容易判断“这个作者大概率属于哪个学院”，所以它们会被用来增强作者归属映射。", first_line=None)
    add_paragraph(doc, "7. 表14的代表作者选择，遵循一个固定顺序：先看这篇论文里有没有作者在 author_map 中明确属于当前学院；如果没有，再看 2024 年种子表里有没有明确归到该学院的人；如果还没有，但已经识别出中文作者，则临时取第一个中文作者；如果连中文作者都没有，则改取第一个可显示拼音。", first_line=None)
    add_paragraph(doc, "8. 表14的 author_counts 就是在上一步基础上累计出来的。循环方式是：对每一篇年度论文，遍历它涉及的每个学院；对每个学院只选1位代表作者，然后记 1 次。因此表14统计的是“学院下由谁代表了多少篇论文”，不是“所有作者人次总和”。", first_line=None)
    add_paragraph(doc, "9. 最后再单独处理“拼音待核”情况。这里又分成两种：一种是整篇论文都没有可靠中文名，此时论文会进入 fully_unresolved 清单；另一种是已经识别出部分中文作者，但仍有少数作者没有中文名，这些论文会进入 pinyin_pending_papers 清单，并在去重总表中把中文作者和待核拼音作者一起展示。", first_line=None)

    add_heading(doc, "四、完整伪代码")
    add_code_block(
        doc,
        [
            "records = []",
            "for 每一期附表 in 6期数据源:",
            "    for 每一行论文记录 in 当前附表:",
            "        record = PaperRecord(...)",
            "        record.author_entities = extract_author_entities(record.authors_raw)",
            "        records.append(record)",
            "",
            "alias_map = build_author_alias_map(records)",
            "for record in records:",
            "    record.chinese_names, record.unresolved_author_tokens = resolve_record_authors(record, alias_map)",
            "",
            "aggregates = {}",
            "for record in records:",
            "    agg = aggregates.setdefault(record.title_key, PaperAggregate(...))",
            "    agg.add(record)",
            "",
            "author_map = infer_author_mapping(aggregates.values())",
            "for agg in aggregates.values():",
            "    for college in agg.colleges:",
            "        author = choose_author_for_college(agg, college, author_map)",
            "        author_counts[college][author] += 1",
            "",
            "for agg in aggregates.values():",
            "    if agg 仍有拼音 token 未核定:",
            "        写入 pinyin_pending_papers",
            "        if agg 一个中文作者都没有:",
            "            同时写入 unresolved_author_papers",
        ],
    )

    add_heading(doc, "五、为什么总表作者列和表14会看起来不一样")
    add_paragraph(
        doc,
        "这是作者统计里最容易混淆的地方。去重总表的“本校作者”列，目的是尽量把一篇论文里已识别出来的本校作者都展示出来，因此可以一篇论文列多个人；而表14的目的是做学院作者分布，因此同一篇论文在同一学院内只允许贡献 1 个作者标签。两者都基于同一套清洗结果，但服务的统计目标不同。",
    )
    add_table(doc, build_examples_rows(), [2000, 2200, 1500, 2260], header_fill="FBE5D6")

    add_heading(doc, "六、2025年当前仍需人工跟进的作者")
    add_paragraph(
        doc,
        "截至当前版本，仍保留拼音待核的论文共有3篇。其中2篇已经识别出部分中文作者，只剩个别作者待核；另有1篇整篇都未能安全核出中文名，只能暂按附表拼音列示。",
    )
    add_table(doc, build_pending_rows(), [2400, 1400, 1800, 2360], header_fill="FFF2CC")

    add_heading(doc, "七、这套方法的优点与边界")
    add_paragraph(doc, "1. 优点在于：规则固定、可以复跑、不会因为人工目测而在总量上反复变化；而且同一作者如果在不同期里出现“拼音+中文”和“只有拼音”两种写法，可以通过别名映射自动并起来。", first_line=None)
    add_paragraph(doc, "2. 它的边界也很明确：程序只能在现有附表、历史种子表和已确认别名的范围内工作。对没有中文括注、又缺少校内公开信息支撑的作者，程序不会硬猜中文名，而是宁可保留拼音待核。", first_line=None)
    add_paragraph(doc, "3. 因此，当前这套作者统计口径适合正式汇报和表格编制；但如果后续需要形成“完整中文作者名录”，仍建议对待核清单逐篇结合原论文全文、学院科研秘书台账或作者本人信息继续补证。", first_line=None)

    doc.save(OUTPUT)
    return OUTPUT


def main() -> None:
    args = parse_args()
    output = generate_author_flow_doc(args.analysis_dir, args.output_docx)
    print(output)


if __name__ == "__main__":
    main()
