from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent
SOURCE_MD = PROJECT_ROOT / "windows_tool" / "技术文档.md"
OUTPUT = PROJECT_ROOT / "outputs" / "Windows版ESI统计工具技术文档.docx"


def ensure_fonts(run, *, ascii_font: str, east_asia_font: str, size: float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:cs"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def add_paragraph(doc: Document, text: str, *, size: float = 12, bold: bool = False, indent: float = 0.74) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Cm(indent)
    run = p.add_run(text)
    ensure_fonts(run, ascii_font="Times New Roman", east_asia_font="仿宋", size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    run = p.add_run(text)
    ensure_fonts(
        run,
        ascii_font="Times New Roman",
        east_asia_font="黑体",
        size=16 if level == 1 else 14 if level == 2 else 12,
        bold=True,
    )


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Cm(0.8)
    fmt.right_indent = Cm(0.4)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    ensure_fonts(run, ascii_font="Courier New", east_asia_font="等线", size=9.5)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="将 Windows 工具技术文档 Markdown 转成 DOCX")
    parser.add_argument("--source-md", type=Path, default=SOURCE_MD, help="Markdown 文档路径")
    parser.add_argument("--output-docx", type=Path, default=OUTPUT, help="输出 DOCX 路径")
    return parser.parse_args()


def generate_windows_tool_tech_doc(source_md: Path, output_docx: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    in_code = False
    lines = source_md.read_text(encoding="utf-8").splitlines()
    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            add_code(doc, line)
            continue
        if not line.strip():
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
            continue
        if line.startswith("- "):
            add_paragraph(doc, "• " + line[2:].strip(), indent=0.4)
            continue
        if line[0].isdigit() and ". " in line[:4]:
            add_paragraph(doc, line, indent=0.4)
            continue
        add_paragraph(doc, line)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    return output_docx


def main() -> None:
    args = parse_args()
    output = generate_windows_tool_tech_doc(args.source_md, args.output_docx)
    print(output)


if __name__ == "__main__":
    main()
